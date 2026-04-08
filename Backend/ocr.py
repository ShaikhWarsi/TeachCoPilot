"""
OCR Module - Local text extraction from PDF, DOCX, and images
Uses pdf2image, pytesseract, and python-docx (NO paid APIs)
"""

import os
import io
import sys
from pathlib import Path
from typing import Optional

# OCR Dependencies
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx import Document


def _validate_tesseract():
    """Validate Tesseract OCR is installed and accessible"""
    try:
        version = pytesseract.get_tesseract_version()
        return True, f"Tesseract v{version} found"
    except Exception as e:
        return False, (
            "Tesseract OCR is not installed or not in PATH.\n"
            "Windows users: Install from https://github.com/UB-Mannheim/tesseract/wiki\n"
            "Then add to PATH or set pytesseract.pytesseract.tesseract_cmd manually\n"
            f"Error: {str(e)}"
        )


def _validate_poppler():
    """Validate poppler is installed (required for PDF processing)"""
    if os.name != 'nt':
        return True, "Poppler check skipped (non-Windows)"
    
    poppler_paths = [
        r'C:\Program Files\poppler\Library\bin',
        r'C:\Program Files (x86)\poppler\Library\bin',
        r'C:\poppler\Library\bin',
    ]
    for path in poppler_paths:
        if os.path.exists(path):
            return True, f"Poppler found at {path}"
    
    return False, (
        "Poppler not found. PDF processing will fail.\n"
        "Windows users: Download from https://github.com/oschwartz10612/poppler-windows/releases/\n"
        "Extract to C:\\poppler and add C:\\poppler\\Library\\bin to PATH"
    )


# Run validation on module load
tesseract_ok, tesseract_msg = _validate_tesseract()
poppler_ok, poppler_msg = _validate_poppler()

if not tesseract_ok:
    print(f"⚠️ OCR Warning: {tesseract_msg}", file=sys.stderr)
if not poppler_ok:
    print(f"⚠️ OCR Warning: {poppler_msg}", file=sys.stderr)

# Configure pytesseract path for Windows
if os.name == 'nt':
    # Common Tesseract installation paths on Windows
    possible_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Users\%USERNAME%\AppData\Local\Tesseract-OCR\tesseract.exe',
    ]
    for path in possible_paths:
        expanded_path = os.path.expandvars(path)
        if os.path.exists(expanded_path):
            pytesseract.pytesseract.tesseract_cmd = expanded_path
            break


class OCRExtractor:
    """Handles text extraction from various file formats"""
    
    @staticmethod
    def extract_text(file_path: str, max_file_size_mb: int = 50) -> str:
        """
        Extract text from file based on its extension
        
        Args:
            file_path: Path to the file
            max_file_size_mb: Maximum file size in MB (default 50)
            
        Returns:
            Extracted text as string
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # Check file size before processing
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > max_file_size_mb:
            raise ValueError(
                f"File too large: {file_size_mb:.1f}MB exceeds maximum of {max_file_size_mb}MB. "
                "Please compress or split the file."
            )
        
        if extension == '.pdf':
            return OCRExtractor._extract_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return OCRExtractor._extract_from_docx(file_path)
        elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
            return OCRExtractor._extract_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    @staticmethod
    def _extract_from_pdf(file_path: Path) -> str:
        """
        Extract text from PDF using pdf2image + pytesseract
        Converts each page to image and runs OCR
        """
        try:
            print(f"Converting PDF to images: {file_path}")
            
            # Convert PDF pages to images
            # poppler_path might be needed on Windows
            poppler_path = None
            if os.name == 'nt':
                # Common poppler paths on Windows
                poppler_paths = [
                    r'C:\Program Files\poppler\Library\bin',
                    r'C:\Program Files (x86)\poppler\Library\bin',
                    r'C:\poppler\Library\bin',
                ]
                for path in poppler_paths:
                    if os.path.exists(path):
                        poppler_path = path
                        break
            
            if poppler_path:
                images = convert_from_path(
                    file_path, 
                    dpi=200,
                    poppler_path=poppler_path
                )
            else:
                images = convert_from_path(file_path, dpi=200)
            
            # Extract text from each page
            all_text = []
            for i, image in enumerate(images):
                print(f"Processing page {i + 1}/{len(images)}...")
                text = pytesseract.image_to_string(image, lang='eng')
                all_text.append(f"--- Page {i + 1} ---\n{text}")
            
            return '\n\n'.join(all_text)
            
        except Exception as e:
            print(f"PDF extraction error: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def _extract_from_docx(file_path: Path) -> str:
        """
        Extract text from DOCX using python-docx
        """
        try:
            print(f"Extracting from DOCX: {file_path}")
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            return '\n\n'.join(paragraphs)
            
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def _extract_from_image(file_path: Path) -> str:
        """
        Extract text from image using pytesseract OCR
        """
        try:
            print(f"Running OCR on image: {file_path}")
            
            # Open and preprocess image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode not in ['RGB', 'L']:
                image = image.convert('RGB')
            
            # Run OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            return text
            
        except Exception as e:
            print(f"Image OCR error: {e}")
            raise Exception(f"Failed to extract text from image: {str(e)}")
    
    @staticmethod
    def cleanup_file(file_path: str) -> None:
        """Remove uploaded file after processing"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Cleaned up: {file_path}")
        except Exception as e:
            print(f"Cleanup error: {e}")


# Convenience function
def extract_text(file_path: str) -> str:
    """Main entry point for text extraction"""
    return OCRExtractor.extract_text(file_path)
